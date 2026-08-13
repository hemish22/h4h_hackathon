"""Bundle the project's source code into a single Word document for submission.
Run: python scripts/make_source_doc.py  ->  Source_Code.docx
"""
import glob
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ordered list of what to include (globs expanded, de-duped, in this order)
GROUPS = [
    ("Project overview", ["README.md"]),
    ("Dependencies", ["requirements.txt", "packages.txt"]),
    ("Core library (src/)", [
        "src/config.py", "src/data_utils.py", "src/audit.py", "src/splits.py",
        "src/match_features.py", "src/build_manifest.py", "src/datasets.py",
        "src/encoders.py", "src/losses.py", "src/fusion.py", "src/decision_fusion.py",
        "src/train.py", "src/evaluate.py", "src/uncertainty.py", "src/explain.py",
        "src/run_explain.py", "src/run_shap.py", "src/eval_extras.py", "src/make_report.py",
    ]),
    ("Prototype (Streamlit app)", ["app/streamlit_app.py"]),
    ("Tests", ["tests/test_smoke.py"]),
    ("Configs", sorted(glob.glob("configs/*.yaml"))),
    ("Scripts", ["scripts/make_diagrams.py", "scripts/make_ppt.py",
                 "scripts/make_report.py" if os.path.exists("scripts/make_report.py") else None,
                 "scripts/fetch_dataset.sh", "scripts/gpu_bootstrap.sh",
                 "scripts/push_dataset.sh", "scripts/gpu_bootstrap.ps1"]),
]

LANG = {".py": "Python", ".yaml": "YAML", ".yml": "YAML", ".sh": "Bash",
        ".ps1": "PowerShell", ".md": "Markdown", ".txt": "Text"}


def add_code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x22, 0x28, 0x31)


def main():
    os.chdir(ROOT)
    doc = Document()
    # normal style tweaks
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

    # title page
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Explainable Multimodal Mental-Health Screening"); r.bold = True; r.font.size = Pt(22)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Source Code — Hack4Hackathon submission").font.size = Pt(13)
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run(f"GitHub: https://github.com/hemish22/h4h_hackathon\nGenerated {date.today().isoformat()}").font.size = Pt(11)
    doc.add_paragraph()

    # table of contents
    doc.add_heading("Contents", level=1)
    files_flat = []
    for group, files in GROUPS:
        files = [f for f in files if f and os.path.exists(f)]
        if not files:
            continue
        doc.add_paragraph(group, style="List Bullet")
        for f in files:
            doc.add_paragraph(f, style="List Bullet 2")
            files_flat.append((group, f))
    doc.add_page_break()

    # each file
    n = 0
    for group, files in GROUPS:
        files = [f for f in files if f and os.path.exists(f)]
        if not files:
            continue
        for f in files:
            n += 1
            ext = os.path.splitext(f)[1]
            doc.add_heading(f, level=2)
            sub = doc.add_paragraph()
            sub.add_run(f"{LANG.get(ext, 'Text')} · {group}").italic = True
            with open(f, encoding="utf-8") as fh:
                add_code(doc, fh.read().rstrip("\n"))
            doc.add_page_break()

    out = os.path.join(ROOT, "Source_Code.docx")
    doc.save(out)
    print(f"wrote {out}  ({n} files)")


if __name__ == "__main__":
    main()
